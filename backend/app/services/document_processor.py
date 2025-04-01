import io
import logging
import os
from pathlib import Path
from typing import Any, Dict, Optional

import docx
import html2text
import PyPDF2
import pytesseract
from PIL import Image

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document formats and converting them to text"""

    # Maximum file size (50MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB in bytes

    SUPPORTED_FORMATS = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".html": "text/html",
        ".htm": "text/html",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".tiff": "image/tiff",
    }

    def __init__(self):
        logger.debug("Initializing DocumentProcessor")
        self.h2t = html2text.HTML2Text()
        self.h2t.ignore_links = True
        self.h2t.ignore_images = True
        self.h2t.ignore_tables = False
        logger.debug("DocumentProcessor initialized")

    async def process_document(
        self,
        file_path: str,
        agent_context: Optional[str] = None,
        original_filename: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Process a document and convert it to text

        Args:
            file_path: Path to the document file
            agent_context: Optional context to be passed to the document analyzer agent
            original_filename: Original filename to use when saving processed text

        Returns:
            Dictionary containing:
                - text: Extracted text content
                - metadata: Document metadata
                - format: Original file format
                - processing_info: Additional processing details
                - agent_context: Context for the document analyzer agent
        """
        try:
            logger.info(f"Starting document processing for file: {file_path}")
            file_path = Path(file_path)

            # Check if file exists
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                raise FileNotFoundError(f"File not found: {file_path}")

            # Check file size
            file_size = file_path.stat().st_size
            logger.info(f"File size: {file_size} bytes")
            if file_size > self.MAX_FILE_SIZE:
                logger.error(
                    f"File size exceeds limit: {file_size} > {self.MAX_FILE_SIZE}"
                )
                raise ValueError(
                    f"File size exceeds maximum limit of {self.MAX_FILE_SIZE / (1024*1024)}MB"
                )

            file_extension = file_path.suffix.lower()
            logger.info(f"Processing file with extension: {file_extension}")

            if file_extension not in self.SUPPORTED_FORMATS:
                logger.error(f"Unsupported file format: {file_extension}")
                raise ValueError(f"Unsupported file format: {file_extension}")

            # Process based on file type
            logger.info(f"Starting {file_extension} processing")
            if file_extension in [".pdf"]:
                logger.debug("Processing PDF file")
                result = await self._process_pdf(file_path)
            elif file_extension in [".docx", ".doc"]:
                logger.debug("Processing Word file")
                result = await self._process_word(file_path)
            elif file_extension in [".txt"]:
                logger.debug("Processing text file")
                result = await self._process_text(file_path)
            elif file_extension in [".html", ".htm"]:
                logger.debug("Processing HTML file")
                result = await self._process_html(file_path)
            elif file_extension in [".png", ".jpg", ".jpeg", ".tiff"]:
                logger.debug("Processing image file")
                result = await self._process_image(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                raise ValueError(f"Unsupported file format: {file_extension}")

            logger.info(
                f"Document processed successfully. Text length: {len(result['text'])}"
            )

            # Save the extracted text to a file if original filename is provided
            if original_filename:
                base_name = os.path.splitext(original_filename)[0]
                await self._save_extracted_text(result["text"], base_name)

            # Add agent context to the result
            if agent_context:
                logger.debug(f"Adding agent context to result: {agent_context}")
                result["agent_context"] = agent_context

            return result

        except Exception as e:
            logger.error(
                f"Error processing document {file_path}: {str(e)}", exc_info=True
            )
            raise

    async def _save_extracted_text(self, text: str, original_filename: str) -> None:
        """
        Save the extracted text to a file in the outputs/anal folder

        Args:
            text: The extracted text content
            original_filename: The original filename (without extension)
        """
        try:
            # Get the backend directory (2 levels up from this file)
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            # Create outputs/anal directory if it doesn't exist
            output_dir = os.path.join(base_dir, "outputs", "anal")
            logger.debug(f"Creating output directory: {output_dir}")
            os.makedirs(output_dir, exist_ok=True)

            # Create the output file path
            output_file = os.path.join(output_dir, f"{original_filename}.txt")
            logger.debug(f"Saving text to: {output_file}")

            # Write the text to the file
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(text)

            logger.info(f"Successfully saved extracted text to {output_file}")
        except Exception as e:
            logger.error(f"Error saving extracted text: {str(e)}", exc_info=True)
            raise

    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF files"""
        text = []
        metadata = {}
        processing_info = {
            "pages": 0,
            "is_scanned": False,
            "ocr_used": False,
            "error": None,
            "status": "processing",
            "current_page": 0,
        }

        try:
            logger.info(f"Starting PDF processing for file: {file_path}")

            # Validate PDF file
            if not file_path.exists():
                processing_info["error"] = "File not found"
                processing_info["status"] = "error"
                logger.error(f"PDF file not found: {file_path}")
                raise FileNotFoundError(f"File not found: {file_path}")

            # Check if file is a valid PDF
            try:
                with open(file_path, "rb") as file:
                    # Read first 4 bytes to check PDF signature
                    header = file.read(4)
                    if header != b"%PDF":
                        processing_info["error"] = "Invalid PDF file"
                        processing_info["status"] = "error"
                        logger.error(f"Invalid PDF signature in file: {file_path}")
                        raise ValueError("Invalid PDF file format")

                    # Reset file pointer
                    file.seek(0)
                    pdf_reader = PyPDF2.PdfReader(file)

                    # Extract metadata
                    metadata = pdf_reader.metadata if pdf_reader.metadata else {}
                    processing_info["pages"] = len(pdf_reader.pages)
                    logger.info(f"PDF has {processing_info['pages']} pages")

                    # Extract text from each page
                    for i, page in enumerate(pdf_reader.pages):
                        try:
                            processing_info["current_page"] = i + 1
                            processing_info[
                                "status"
                            ] = f"Processing page {i+1}/{len(pdf_reader.pages)}"
                            logger.debug(
                                f"Processing page {i+1}/{len(pdf_reader.pages)}"
                            )

                            page_text = page.extract_text()
                            if page_text.strip():
                                text.append(page_text)
                                logger.debug(
                                    f"Successfully extracted text from page {i+1}, length: {len(page_text)}"
                                )
                            else:
                                logger.warning(f"No text extracted from page {i+1}")
                                text.append("")
                        except Exception as page_error:
                            logger.warning(
                                f"Error extracting text from page {i+1}: {str(page_error)}"
                            )
                            text.append("")  # Add empty string for failed page
                            continue

                    # Handle scanned PDFs if no text was extracted
                    if not any(text):
                        logger.warning(
                            f"No text extracted from PDF {file_path}, attempting OCR"
                        )
                        processing_info["is_scanned"] = True
                        processing_info["status"] = "Running OCR"
                        text = await self._process_scanned_pdf(file_path)
                        processing_info["ocr_used"] = True
                        logger.info("OCR completed successfully")

            except Exception as pdf_error:
                processing_info["error"] = str(pdf_error)
                processing_info["status"] = "error"
                logger.error(f"Error processing PDF: {str(pdf_error)}", exc_info=True)
                raise

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}", exc_info=True)
            processing_info["error"] = str(e)
            processing_info["status"] = "error"
            raise

        processing_info["status"] = "completed"
        logger.info(
            f"PDF processing completed successfully. Total text length: {len(''.join(text))}"
        )
        return {
            "text": "\n".join(text),
            "metadata": metadata,
            "format": "pdf",
            "processing_info": processing_info,
        }

    async def _process_scanned_pdf(self, file_path: Path) -> list[str]:
        """Process scanned PDFs using OCR"""
        try:
            # Convert PDF pages to images
            images = await self._pdf_to_images(file_path)

            # Process each image with OCR
            text = []
            for image in images:
                text.append(pytesseract.image_to_string(image))

            return text

        except Exception as e:
            logger.error(f"Error processing scanned PDF {file_path}: {str(e)}")
            raise

    async def _pdf_to_images(self, file_path: Path) -> list[Image.Image]:
        """Convert PDF pages to PIL Images"""
        try:
            from pdf2image import convert_from_path

            return convert_from_path(file_path)
        except ImportError:
            logger.error(
                "pdf2image not installed. Please install it with: pip install pdf2image"
            )
            raise ImportError(
                "pdf2image package is required for processing scanned PDFs"
            )

    async def _process_word(self, file_path: Path) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            doc = docx.Document(file_path)
            text = [paragraph.text for paragraph in doc.paragraphs]

            # Extract metadata
            metadata = {
                "author": doc.core_properties.author,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
                "title": doc.core_properties.title,
            }

        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            raise

        return {
            "text": "\n".join(text),
            "metadata": metadata,
            "format": file_path.suffix.lower(),
            "processing_info": {"paragraphs": len(doc.paragraphs)},
        }

    async def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            logger.debug(f"Processing text file: {file_path}")
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()
                logger.debug(f"Text file content length: {len(text)}")

        except Exception as e:
            logger.error(
                f"Error processing text file {file_path}: {str(e)}", exc_info=True
            )
            raise

        return {
            "text": text,
            "metadata": {},
            "format": "txt",
            "processing_info": {"lines": text.count("\n") + 1},
        }

    async def _process_html(self, file_path: Path) -> Dict[str, Any]:
        """Process HTML files"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                html_content = file.read()

            # Convert HTML to text
            text = self.h2t.handle(html_content)

        except Exception as e:
            logger.error(f"Error processing HTML file {file_path}: {str(e)}")
            raise

        return {
            "text": text,
            "metadata": {},
            "format": "html",
            "processing_info": {"links_ignored": True, "images_ignored": True},
        }

    async def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process image files using OCR"""
        try:
            # Open image
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image)

            # Get image metadata
            metadata = {"format": image.format, "size": image.size, "mode": image.mode}

        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}")
            raise

        return {
            "text": text,
            "metadata": metadata,
            "format": file_path.suffix.lower(),
            "processing_info": {
                "ocr_used": True,
                "confidence": "medium",  # TODO: Implement confidence scoring
            },
        }
